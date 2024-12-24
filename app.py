from flask import Flask, render_template, request, jsonify
import os
from docx import Document

app = Flask(__name__)

# Configuración
app.config['UPLOAD_FOLDER'] = 'plantillas_actas/'
app.config['DILIGENCIAS_FOLDER'] = 'diligencias/'
app.config['EXPORT_FOLDER'] = 'export/'

# Crear carpetas necesarias si no existen
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DILIGENCIAS_FOLDER'], exist_ok=True)
os.makedirs(app.config['EXPORT_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    # Cargar las plantillas disponibles en la carpeta
    plantillas = os.listdir(app.config['UPLOAD_FOLDER'])
    return render_template('index.html', plantillas=plantillas)

@app.route('/crear_acta', methods=['POST'])
def crear_acta():
    nombre_diligencia = request.form['nombre_diligencia']
    plantilla_seleccionada = request.form['plantilla']

    # Crear carpeta para la diligencia si no existe
    ruta_diligencia = os.path.join(app.config['DILIGENCIAS_FOLDER'], nombre_diligencia)
    os.makedirs(ruta_diligencia, exist_ok=True)

    # Copiar la plantilla seleccionada a la carpeta de la diligencia
    plantilla_path = os.path.join(app.config['UPLOAD_FOLDER'], plantilla_seleccionada)
    nuevo_acta_path = os.path.join(ruta_diligencia, f"Acta_{plantilla_seleccionada}")

    # Generar un nuevo archivo Word basado en la plantilla
    doc = Document(plantilla_path)
    # Aquí puedes modificar el contenido del documento si es necesario
    doc.save(nuevo_acta_path)

    return jsonify({"message": "Acta creada exitosamente", "ruta": nuevo_acta_path})

@app.route('/exportar', methods=['POST'])
def exportar():
    # Obtener datos del cuerpo JSON
    data = request.json
    if not data or 'nombre_diligencia' not in data:
        return jsonify({"error": "El campo 'nombre_diligencia' es requerido."}), 400

    nombre_diligencia = data['nombre_diligencia']
    ruta_diligencia = os.path.join(app.config['DILIGENCIAS_FOLDER'], nombre_diligencia)
    ruta_export = os.path.join(app.config['EXPORT_FOLDER'], nombre_diligencia)

    # Copiar todas las actas creadas a la carpeta de exportación
    os.makedirs(ruta_export, exist_ok=True)
    for acta in os.listdir(ruta_diligencia):
        origen = os.path.join(ruta_diligencia, acta)
        destino = os.path.join(ruta_export, acta)
        with open(origen, 'rb') as f_origen:
            with open(destino, 'wb') as f_destino:
                f_destino.write(f_origen.read())

    return jsonify({"message": "Actas exportadas exitosamente", "ruta": ruta_export})

@app.route('/plantillas', methods=['GET'])
def obtener_plantillas():
    plantillas = os.listdir(app.config['UPLOAD_FOLDER'])
    return jsonify(plantillas)

@app.route('/guardar_acta', methods=['POST'])
def guardar_acta():
    data = request.json
    plantilla = data['plantilla']
    nombre_diligencia = data['nombreDiligencia']
    datos = data['datos']

    # Crear carpeta para la diligencia
    ruta_diligencia = os.path.join(app.config['DILIGENCIAS_FOLDER'], nombre_diligencia)
    os.makedirs(ruta_diligencia, exist_ok=True)

    # Crear el acta basada en la plantilla
    plantilla_path = os.path.join(app.config['UPLOAD_FOLDER'], plantilla)
    nuevo_acta_path = os.path.join(ruta_diligencia, f"{datos['nombre']}.docx")

    doc = Document(plantilla_path)
    for key, value in datos.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(nuevo_acta_path)

    return jsonify({"message": "Acta guardada exitosamente."})

if __name__ == '__main__':
    app.run(debug=True)
